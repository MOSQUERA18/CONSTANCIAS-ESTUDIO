import os
import re
import pandas as pd
from docx import Document
from datetime import datetime
import logging
from pathlib import Path
import tkinter as tk
from tkinter import filedialog, messagebox
import shutil

# Configurar logging
logging.basicConfig(level=logging.INFO, format='[%(levelname)s] %(message)s')

# Rutas del sistema
DOCUMENTS_PATH = Path.home() / "Documents"
DOWNLOADS_PATH = Path.home() / "Downloads"
OUTPUT_DIR = os.path.join(DOCUMENTS_PATH, "Certificados-Constancias")

# Archivos
BASE_DIR = Path(__file__).resolve().parent
TEMPLATE_FILE = BASE_DIR / "plantilla_constancias.docx"
PLANTILLA_EXCEL_ORIGINAL = "plantilla_constancias.xlsx"

# Crear carpeta de salida si no existe
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Funciones utilitarias
def formatear_fecha_hoy():
    hoy = datetime.today()
    meses_es = [
        "enero", "febrero", "marzo", "abril", "mayo", "junio",
        "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"
    ]
    return f"a los {hoy.day} días del mes de {meses_es[hoy.month - 1]}"

def reemplazar_en_parrafos(parrafos, reemplazos):
    for p in parrafos:
        texto_original = p.text
        texto_reemplazado = texto_original
        for clave, valor in reemplazos.items():
            texto_reemplazado = texto_reemplazado.replace(f"{{{{{clave}}}}}", str(valor))
        if texto_reemplazado != texto_original:
            p.clear()  # Limpia el párrafo completo correctamente
            p.add_run(texto_reemplazado)


def reemplazar_campos(doc, reemplazos):
    reemplazar_en_parrafos(doc.paragraphs, reemplazos)


def obtener_campos_docx(doc):
    texto = "\n".join([p.text for p in doc.paragraphs])
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texto += "\n" + cell.text
    campos = set(re.findall(r"{{(.*?)}}", texto))
    return campos

def validar_columnas(df, columnas_requeridas):
    faltantes = [col for col in columnas_requeridas if col not in df.columns]
    if faltantes:
        raise ValueError(f"Columnas faltantes en el archivo Excel: {faltantes}")

def formatear_fecha_excel(fecha_excel):
    if pd.isnull(fecha_excel):
        return ""
    
    meses_es = {
        1: "enero", 2: "febrero", 3: "marzo", 4: "abril",
        5: "mayo", 6: "junio", 7: "julio", 8: "agosto",
        9: "septiembre", 10: "octubre", 11: "noviembre", 12: "diciembre"
    }
    
    return f"{fecha_excel.day} de {meses_es[fecha_excel.month]} de {fecha_excel.year}"



# Función para crear y descargar la plantilla de Excel desde cero
def descargar_plantilla():
    try:
        columnas = [
            "nombre", "tipo", "id", "programa", "ficha",
            "fecha1", "fecha2", "fecha3", "fecha4", "interesar"
        ]
        
        # Crear DataFrame vacío con encabezados
        df = pd.DataFrame(columns=columnas)

        # Ruta de destino
        destino = os.path.join(DOWNLOADS_PATH, "plantilla_constancias.xlsx")
        df.to_excel(destino, index=False)

        messagebox.showinfo("Descarga completada", f"Plantilla Excel creada en:\n{destino}")
    except Exception as e:
        messagebox.showerror("Error", f"No se pudo crear la plantilla:\n{e}")



# Función principal para generar certificados
def generar_certificados_desde_excel(excel_path):
    try:
        df = pd.read_excel(excel_path)
        df.columns = df.columns.str.strip().str.replace(" ", "_").str.lower()
        columnas_requeridas = ["nombre", "tipo", "id", "programa", "ficha", "fecha1", "fecha2", "fecha3", "fecha4", "interesar"]
        validar_columnas(df, columnas_requeridas)

        fecha_actual_formateada = formatear_fecha_hoy()
        total = 0


        for _, fila in df.iterrows():
            try:
                doc = Document(TEMPLATE_FILE)

                reemplazos = {
                    "NOMBRE": fila["nombre"],
                    "TIPO_DE_DOCUMENTO": fila["tipo"],
                    "NUMERO_IDENTIFICACION": fila["id"],
                    "NOMBRE DEL PROGRAMA": fila["programa"],
                    "NUMERO_FICHA": fila["ficha"],
                    "FECHA_ONE": formatear_fecha_excel(fila["fecha1"]),
                    "FECHA_TWO": formatear_fecha_excel(fila["fecha2"]),
                    "FECHA_THREE": formatear_fecha_excel(fila["fecha3"]),
                    "FECHA_FOUR": formatear_fecha_excel(fila["fecha4"]),
                    "interesar": fila["interesar"],
                    "DIA_REALIZA": fecha_actual_formateada,
                }

                campos_docx = obtener_campos_docx(doc)
                for campo in campos_docx:
                    if campo not in reemplazos:
                        logging.warning(f"El campo {{{{{campo}}}}} no tiene un valor definido.")

                reemplazar_campos(doc, reemplazos)

                nombre_limpio = re.sub(r"[^a-zA-Z0-9áéíóúÁÉÍÓÚñÑ_\-]", "_", fila["nombre"])
                nombre_archivo = f"CONSTANCIA_{nombre_limpio}.docx"
                doc.save(os.path.join(OUTPUT_DIR, nombre_archivo))
                total += 1

            except Exception as e:
                logging.error(f"❌ Error procesando la fila: {fila.to_dict()}\n{e}")

        messagebox.showinfo("Proceso finalizado", f"{total} certificados generados en:\n{OUTPUT_DIR}")

    except Exception as e:
        messagebox.showerror("Error", str(e))

# Función para seleccionar el archivo Excel
def seleccionar_archivo():
    filepath = filedialog.askopenfilename(
        filetypes=[("Archivos Excel", "*.xlsx *.xls")]
    )
    if filepath:
        generar_certificados_desde_excel(filepath)

# Interfaz gráfica
def iniciar_app():
    ventana = tk.Tk()
    ventana.title("Generador de Certificados")
    ventana.geometry("400x250")
    ventana.resizable(False, False)

    tk.Label(ventana, text="Generador de Constancias", font=("Arial", 16)).pack(pady=20)

    tk.Button(ventana, text="📥 Descargar plantilla Excel", width=30, command=descargar_plantilla).pack(pady=10)
    tk.Button(ventana, text="📂 Seleccionar archivo Excel", width=30, command=seleccionar_archivo).pack(pady=10)

    tk.Label(ventana, text=f"Los certificados se guardarán en:\n{OUTPUT_DIR}", font=("Arial", 10)).pack(pady=20)

    ventana.mainloop()

if __name__ == "__main__":
    iniciar_app()